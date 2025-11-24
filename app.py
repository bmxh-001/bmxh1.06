import requests
import os
import uuid
import sys
from flask import Flask, request, jsonify, render_template, Response, stream_with_context
import re
import tempfile
import json
import logging
from docx import Document
from docx.shared import Inches
import datetime
import io
import warnings # Added for API key warning
import re
import tempfile
import base64
import networkx as nx
import matplotlib.pyplot as plt
from io import BytesIO
import mimetypes

# 添加对PyInstaller打包的支持
def resource_path(relative_path):
    """获取资源的绝对路径，支持开发环境和PyInstaller打包后的环境"""
    try:
        # PyInstaller创建临时文件夹并将路径存储在_MEIPASS中
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")
    
    return os.path.join(base_path, relative_path)

app = Flask(__name__, 
            template_folder=resource_path('templates'),
            static_folder=resource_path('static'))

@app.after_request
def set_js_mime_type(response):
    if request.path.endswith('.js'):
        response.headers['Content-Type'] = 'application/javascript'
    return response

# 设置日志
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

# Configuration
OPENROUTER_API_BASE = "https://openrouter.ai/api/v1" # OpenRouter Base URL
OPENROUTER_API_KEY = None # Will be set via API endpoint
CUSTOM_API_CONFIGS = {}  # Will store custom API configurations in memory

# 确定是否在打包环境中运行
if getattr(sys, 'frozen', False):
    # 如果是打包后的EXE，使用应用程序所在目录
    application_path = os.path.dirname(sys.executable)
    EXPORT_DIR = os.path.join(application_path, 'exports')
    DATA_DIR = os.path.join(application_path, 'data')  # 数据文件目录
    MODEL_DIR = os.path.join(application_path, 'model') # 新增：模型模板文件夹
else:
    # 正常开发环境
    EXPORT_DIR = os.path.join(os.path.dirname(__file__), 'exports')
    DATA_DIR = os.path.join(os.path.dirname(__file__), 'data')  # 数据文件目录
    MODEL_DIR = os.path.join(os.path.dirname(__file__), 'model') # 新增：模型模板文件夹

# 数据文件路径
NOVELS_FILE = os.path.join(DATA_DIR, 'novels.json')
CHARACTERS_FILE = os.path.join(DATA_DIR, 'characters.json')
GLOSSARY_FILE = os.path.join(DATA_DIR, 'glossary.json')
STYLES_FILE = os.path.join(DATA_DIR, 'styles.json')
GAMES_FILE = os.path.join(DATA_DIR, 'games.json')
API_CONFIGS_FILE = os.path.join(DATA_DIR, 'api_configs.json')  # 自定义API配置文件
API_KEY_FILE = os.path.join(DATA_DIR, 'api_key.json')  # OpenRouter API密钥文件

# Simple in-memory storage for novels (replace with a database for persistence)
novels_db = {}
# Simple in-memory storage for character library
characters_db = {}
# Simple in-memory storage for glossary entries
glossary_db = {}
# Simple in-memory storage for style library
styles_db = {}
# Simple in-memory storage for games
games_db = {}

# Ensure export and data directories exist
if not os.path.exists(EXPORT_DIR):
    os.makedirs(EXPORT_DIR)
if not os.path.exists(DATA_DIR):
    os.makedirs(DATA_DIR)
if not os.path.exists(MODEL_DIR): # 新增：确保模型文件夹存在
    os.makedirs(MODEL_DIR)

# 加载持久化数据
def load_data():
    """从JSON文件加载数据到内存"""
    global novels_db, characters_db, glossary_db, styles_db, games_db, CUSTOM_API_CONFIGS, OPENROUTER_API_KEY
    
    logging.info("尝试加载持久化数据...")
    logging.info(f"小说数据文件路径: {NOVELS_FILE}")

    # 确保所有数据目录存在
    for directory in [DATA_DIR, EXPORT_DIR, MODEL_DIR]:
        if not os.path.exists(directory):
            os.makedirs(directory)
            logging.info(f"创建目录: {directory}")

    # 确保所有数据文件存在
    data_files = {
        NOVELS_FILE: {},
        CHARACTERS_FILE: {},
        GLOSSARY_FILE: {},
        STYLES_FILE: {},
        GAMES_FILE: {},
        API_CONFIGS_FILE: {},
        API_KEY_FILE: {"api_key": None}
    }

    for file_path, default_content in data_files.items():
        if not os.path.exists(file_path):
            try:
                with open(file_path, 'w', encoding='utf-8') as f:
                    json.dump(default_content, f, ensure_ascii=False, indent=2)
                logging.info(f"创建数据文件: {file_path}")
            except Exception as e:
                logging.error(f"创建数据文件失败 {file_path}: {e}")

    # 加载小说数据
    if os.path.exists(NOVELS_FILE):
        logging.info("小说数据文件存在，尝试加载...")
        try:
            with open(NOVELS_FILE, 'r', encoding='utf-8') as f:
                novels_db = json.load(f)
            logging.info(f"成功加载了 {len(novels_db)} 部小说。")
            # 添加日志，打印加载后的小说ID列表和总数
            logging.info(f"加载后的小说ID列表 (前5个): {list(novels_db.keys())[:5]}")
            logging.info(f"加载后的小说总数: {len(novels_db)}")
        except json.JSONDecodeError as e:
            logging.error(f"加载小说数据出错 (JSON解析错误): {e}")
            novels_db = {}
        except Exception as e:
            logging.error(f"加载小说数据出错 (其他错误): {e}")
            novels_db = {}
    else:
        logging.info("小说数据文件不存在，初始化为空数据库。")
        novels_db = {}
       
       
    # 加载角色数据
    if os.path.exists(CHARACTERS_FILE):
        try:
            with open(CHARACTERS_FILE, 'r', encoding='utf-8') as f:
                characters_db = json.load(f)
            logging.info(f"成功加载了 {len(characters_db)} 个角色。")
        except json.JSONDecodeError as e:
            logging.error(f"加载角色数据出错 (JSON解析错误): {e}")
            characters_db = {}
        except Exception as e:
            logging.error(f"加载角色数据出错 (其他错误): {e}")
            characters_db = {}
    else:
        logging.info("角色数据文件不存在，初始化为空数据库。")
        characters_db = {}
    
    # 加载词条数据
    if os.path.exists(GLOSSARY_FILE):
        try:
            with open(GLOSSARY_FILE, 'r', encoding='utf-8') as f:
                glossary_db = json.load(f)
            logging.info(f"成功加载了 {len(glossary_db)} 个词条。")
        except json.JSONDecodeError as e:
            logging.error(f"加载词条数据出错 (JSON解析错误): {e}")
            glossary_db = {}
        except Exception as e:
            logging.error(f"加载词条数据出错 (其他错误): {e}")
            glossary_db = {}
    else:
        logging.info("词条数据文件不存在，初始化为空数据库。")
        glossary_db = {}
            
    # 加载风格数据
    if os.path.exists(STYLES_FILE):
        try:
            with open(STYLES_FILE, 'r', encoding='utf-8') as f:
                styles_db = json.load(f)
            logging.info(f"成功加载了 {len(styles_db)} 个风格词条。")
        except json.JSONDecodeError as e:
            logging.error(f"加载风格数据出错 (JSON解析错误): {e}")
            styles_db = {}
        except Exception as e:
            logging.error(f"加载风格数据出错 (其他错误): {e}")
            styles_db = {}
    else:
        logging.info("风格数据文件不存在，初始化为空数据库。")
        styles_db = {}
      
    # 加载游戏数据
    if os.path.exists(GAMES_FILE):
        try:
            with open(GAMES_FILE, 'r', encoding='utf-8') as f:
                games_db = json.load(f)
            logging.info(f"成功加载了 {len(games_db)} 个游戏记录。")
        except json.JSONDecodeError as e:
            logging.error(f"加载游戏数据出错 (JSON解析错误): {e}")
            games_db = {}
        except Exception as e:
            logging.error(f"加载游戏数据出错 (其他错误): {e}")
            games_db = {}
    else:
        logging.info("游戏数据文件不存在，初始化为空数据库。")
        games_db = {}
            
    # 加载自定义API配置
    if os.path.exists(API_CONFIGS_FILE):
        try:
            with open(API_CONFIGS_FILE, 'r', encoding='utf-8') as f:
                CUSTOM_API_CONFIGS = json.load(f)
            logging.info(f"成功加载了 {len(CUSTOM_API_CONFIGS)} 个自定义API配置。")
        except json.JSONDecodeError as e:
            logging.error(f"加载自定义API配置出错 (JSON解析错误): {e}")
            CUSTOM_API_CONFIGS = {}
        except Exception as e:
            logging.error(f"加载自定义API配置出错 (其他错误): {e}")
            CUSTOM_API_CONFIGS = {}
    else:
        logging.info("自定义API配置文件不存在，初始化为空配置。")
        CUSTOM_API_CONFIGS = {}
            
    # 加载OpenRouter API密钥
    if os.path.exists(API_KEY_FILE):
        try:
            with open(API_KEY_FILE, 'r', encoding='utf-8') as f:
                api_key_data = json.load(f)
                OPENROUTER_API_KEY = api_key_data.get("api_key")
            if OPENROUTER_API_KEY:
                logging.info("成功加载了API密钥。")
        except json.JSONDecodeError as e:
            logging.error(f"加载API密钥出错 (JSON解析错误): {e}")
            OPENROUTER_API_KEY = None
        except Exception as e:
            logging.error(f"加载API密钥出错 (其他错误): {e}")
            OPENROUTER_API_KEY = None
    else:
        print("API密钥文件不存在或密钥未设置。")
        OPENROUTER_API_KEY = None

# 保存数据到文件
def save_novels():
    """保存小说数据到文件"""
    logging.info("尝试保存小说数据...")
    logging.info(f"小说数据文件路径: {NOVELS_FILE}")
    try:
        # 使用临时文件进行原子性写入
        with tempfile.NamedTemporaryFile('w', encoding='utf-8', delete=False, dir=os.path.dirname(NOVELS_FILE)) as tmp_file:
            json.dump(novels_db, tmp_file, ensure_ascii=False, indent=2)
        # 原子性替换原文件
        os.replace(tmp_file.name, NOVELS_FILE)
        logging.info(f"成功保存了 {len(novels_db)} 部小说。")
        return True
    except Exception as e:
        logging.error(f"保存小说数据出错: {e}")
        # 如果发生错误，尝试清理临时文件
        if 'tmp_file' in locals() and os.path.exists(tmp_file.name):
            os.remove(tmp_file.name)
        return False

def save_characters():
    """保存角色数据到文件"""
    try:
        with open(CHARACTERS_FILE, 'w', encoding='utf-8') as f:
            json.dump(characters_db, f, ensure_ascii=False, indent=2)
        print(f"成功保存了 {len(characters_db)} 个角色。")
        return True
    except Exception as e:
        print(f"保存角色数据出错: {e}")
        return False

def save_glossary():
    """保存词条数据到文件"""
    try:
        with open(GLOSSARY_FILE, 'w', encoding='utf-8') as f:
            json.dump(glossary_db, f, ensure_ascii=False, indent=2)
        print(f"成功保存了 {len(glossary_db)} 个词条。")
        return True
    except Exception as e:
        print(f"保存词条数据出错: {e}")
        return False

def save_styles():
    """保存风格数据到文件"""
    try:
        with open(STYLES_FILE, 'w', encoding='utf-8') as f:
            json.dump(styles_db, f, ensure_ascii=False, indent=2)
        print(f"成功保存了 {len(styles_db)} 个风格词条。")
        return True
    except Exception as e:
        print(f"保存风格数据出错: {e}")
        return False

def save_games():
    """保存游戏数据到文件"""
    try:
        with open(GAMES_FILE, 'w', encoding='utf-8') as f:
            json.dump(games_db, f, ensure_ascii=False, indent=2)
        print(f"成功保存了 {len(games_db)} 个游戏记录。")
        return True
    except Exception as e:
        print(f"保存游戏数据出错: {e}")
        return False

def save_api_configs():
    """保存自定义API配置到文件"""
    try:
        with open(API_CONFIGS_FILE, 'w', encoding='utf-8') as f:
            json.dump(CUSTOM_API_CONFIGS, f, ensure_ascii=False, indent=2)
        print(f"成功保存了 {len(CUSTOM_API_CONFIGS)} 个自定义API配置。")
        return True
    except Exception as e:
        print(f"保存自定义API配置出错: {e}")
        return False

def save_api_key():
    """保存OpenRouter API密钥到文件"""
    try:
        with open(API_KEY_FILE, 'w', encoding='utf-8') as f:
            json.dump({"api_key": OPENROUTER_API_KEY}, f, ensure_ascii=False, indent=2)
        print("成功保存了API密钥。")
        return True
    except Exception as e:
        print(f"保存API密钥出错: {e}")
        return False

# 在应用启动时加载数据
load_data()

# --- Helper Functions ---

# Renamed function and updated to return specified OpenRouter models
def get_available_models():
    """Returns a list of available models (currently hardcoded for OpenRouter)."""
    # In the future, this could potentially fetch from OpenRouter, but for now,
    # we list the specified free models.
    return [
        {"id": "deepseek/deepseek-chat-v3-0324:free", "name": "Deepseek/deepseek-chat-v3-0324:free (Free)"},
        {"id": "qwen/qwen3-235b-a22b:free", "name": "Qwen/qwen3-235b-a22b:free (Free)"},
        {"id": "google/gemma-3-27b-it:free", "name": "Google Gemma 3 27b IT (Free)"},
        {"id": "qwen/qwen2.5-vl-72b-instruct:free", "name": "qwen/qwen2.5-vl-72b-instruct:free"},
        {"id": "meta-llama/llama-4-scout:free", "name": "Meta-llama/llama-4-scout:free (Free)"},
        {"id": "google/gemini-2.0-flash-exp:free", "name": "Google/gemini-2.0-flash-exp:free (Free)"},
        {"id": "deepseek/deepseek-r1:free", "name": "DeepSeek Coder R1 (Free)"}
    ]
    # --- Removed Ollama fetching logic ---
    # try:
    #     response = requests.get(f"{OLLAMA_BASE_URL}/api/tags")
    #     response.raise_for_status() # Raise an exception for bad status codes (4xx or 5xx)
    #     models_data = response.json()
    #     # Extract just the model names
    #     model_names = [model['name'] for model in models_data.get('models', [])]
    #     return model_names
    # except requests.exceptions.RequestException as e:
    #     print(f"Error connecting to Ollama: {e}")
    #     return None # Indicate failure
    # except json.JSONDecodeError:
    #     print(f"Error decoding JSON response from Ollama: {response.text}")
    #     return None
    # except Exception as e:
    #     print(f"An unexpected error occurred while fetching models: {e}")
    #     return None

# --- API Endpoints ---

@app.route('/')
def index():
    """Serves the main HTML page."""
    return render_template('index.html')

@app.route('/api/models', methods=['GET'])
def list_models():
    """API endpoint to get the list of available models."""
    # Get OpenRouter models
    openrouter_models = get_available_models()
    
    # Get custom API models from configurations
    custom_models = []
    for api_id, config in CUSTOM_API_CONFIGS.items():
        custom_models.append({
            "id": f"custom-{api_id}", 
            "name": f"{config['name']} - {config['modelName']}"
        })
    
    # Combine both lists
    all_models = openrouter_models + custom_models
    
    if all_models:
        return jsonify(all_models)
    else:
        return jsonify({"error": "Could not retrieve model list."}), 500

@app.route('/api/novels', methods=['GET'])
def get_novels():
    """API endpoint to list all novels."""
    # Return a list of novels with id and title
    novel_list = [{"id": novel_id, "title": data.get("title", "Untitled")} for novel_id, data in novels_db.items()]
    return jsonify(novel_list)

@app.route('/api/novels', methods=['POST'])
def create_novel():
    """API endpoint to create a new novel."""
    data = request.get_json()
    if not data or 'title' not in data:
        return jsonify({"message": "Title is required"}), 400

    novel_id = str(uuid.uuid4()) # Generate a unique ID
    new_novel = {
        "title": data['title'],
        "content": data.get('content', ''), # Initialize with empty content or provided content
        "characters": "", # Initialize text field
        "knowledge": "",
        "style_prompt": "", # Initialize text field
        "character_tags": [], # Initialize tag arrays
        "glossary_tags": [],
        "style_tags": []
        # Add other metadata later (e.g., settings, characters)
    }
    novels_db[novel_id] = new_novel
    logging.info(f"Created novel in memory: {novel_id} - {data['title']}")
    
    if save_novels():  # 保存小说数据并检查结果
        logging.info(f"Successfully saved novel {novel_id} to file.")
        return jsonify({"id": novel_id, "title": new_novel["title"]}), 201 # 201 Created status
    else:
        logging.error(f"Failed to save novel {novel_id} to file.")
        # Decide how to handle save failure - maybe return a 500 error?
        # For now, we'll still return 201 as the novel is in memory, but logging indicates the issue.
        # A more robust solution might require rolling back the in-memory addition or returning an error.
        # Given the current structure, logging the error is the most direct fix based on the proposal.
        return jsonify({"id": novel_id, "title": new_novel["title"], "warning": "Failed to save novel to file"}), 201 # Still return 201 but add a warning

@app.route('/api/novels/<novel_id>', methods=['GET'])
def get_novel(novel_id):
    """API endpoint to get details of a specific novel."""
    novel = novels_db.get(novel_id)
    if novel:
        return jsonify({
            "id": novel_id,
            "title": novel.get("title", "Untitled"),
            "content": novel.get("content", ""),
            "characters": novel.get("characters", ""), # Return text field
            "knowledge": novel.get("knowledge", ""),
            "style_prompt": novel.get("style_prompt", ""), # Return text field
            "character_tags": novel.get("character_tags", []), # Return tag arrays
            "glossary_tags": novel.get("glossary_tags", []),
            "style_tags": novel.get("style_tags", [])
            # Add other fields as needed
        })
    else:
        return jsonify({"message": "Novel not found"}), 404

@app.route('/api/novels/<novel_id>', methods=['PUT', 'PATCH'])
def update_novel(novel_id):
    """API endpoint to update a novel's content (or other fields)."""
    if novel_id not in novels_db:
        return jsonify({"message": "Novel not found"}), 404

    data = request.get_json()
    if not data:
        return jsonify({"message": "No update data provided"}), 400

    # Update fields provided in the request
    if 'title' in data:
        novels_db[novel_id]['title'] = data['title']
    if 'content' in data:
        novels_db[novel_id]['content'] = data['content']
    # Add updates for metadata text fields
    if 'characters' in data:
        novels_db[novel_id]['characters'] = data['characters']
    if 'knowledge' in data:
        novels_db[novel_id]['knowledge'] = data['knowledge']
    if 'style_prompt' in data:
        novels_db[novel_id]['style_prompt'] = data['style_prompt']
    # Add updates for tag arrays
    if 'character_tags' in data:
        novels_db[novel_id]['character_tags'] = data['character_tags']
    if 'glossary_tags' in data:
        novels_db[novel_id]['glossary_tags'] = data['glossary_tags']
    if 'style_tags' in data:
        novels_db[novel_id]['style_tags'] = data['style_tags']

    print(f"Updated novel: {novel_id}")
    save_novels()  # 保存小说数据
    return jsonify({"message": "Novel updated successfully", "id": novel_id})

@app.route('/api/novels/<novel_id>', methods=['DELETE'])
def delete_novel(novel_id):
    """API endpoint to delete a specific novel."""
    if novel_id in novels_db:
        deleted_title = novels_db[novel_id].get('title', 'Untitled')
        del novels_db[novel_id]
        print(f"Deleted novel: {novel_id} - {deleted_title}")
        save_novels()  # 保存小说数据
        return jsonify({"message": f"Novel '{deleted_title}' deleted successfully"})
    else:
        print(f"Attempted to delete non-existent novel: {novel_id}")
        return jsonify({"message": "Novel not found"}), 404

@app.route('/api/set-api-key', methods=['POST'])
def set_api_key():
    """API endpoint to set and validate the OpenRouter API key."""
    global OPENROUTER_API_KEY
    data = request.get_json()
    if not data or 'api_key' not in data:
        return jsonify({"error": "请输入API密钥"}), 400
    
    api_key = data['api_key'].strip()
    if not api_key:
        return jsonify({"error": "API密钥不能为空"}), 400
    
    # 验证API密钥格式
    if not api_key.startswith('sk-'):
        return jsonify({"error": "无效的API密钥格式，OpenRouter API密钥应以'sk-'开头"}), 400
    
    # 设置API密钥
    OPENROUTER_API_KEY = api_key
    
    # 保存API密钥到文件
    save_api_key()
    
    return jsonify({"message": "API密钥已成功设置并保存"})

@app.route('/api/api-key-status', methods=['GET'])
def get_api_key_status():
    """API endpoint to check if an API key is set."""
    has_key = OPENROUTER_API_KEY is not None and len(str(OPENROUTER_API_KEY)) > 0
    
    key_preview = None
    if has_key and len(str(OPENROUTER_API_KEY)) >= 11:  # 确保密钥长度足够生成预览
        key_preview = f"{OPENROUTER_API_KEY[:7]}...{OPENROUTER_API_KEY[-4:]}"
    
    return jsonify({
        "has_key": has_key,
        "key_preview": key_preview
    })

@app.route('/api/generate', methods=['POST'])
def generate_text():
    """API endpoint to generate text using OpenRouter or custom APIs."""
    data = request.get_json()
    if not data or 'model' not in data or 'prompt' not in data:
        return jsonify({"error": "Missing required fields: model and prompt"}), 400

    model = data['model']
    prompt = data['prompt']
    temperature = data.get('temperature', 0.7)  # Default temperature to 0.7

    # Check if this is a custom API model
    if model.startswith('custom-'):
        api_id = model.replace('custom-', '', 1)
        if api_id not in CUSTOM_API_CONFIGS:
            return jsonify({"error": "Custom API configuration not found"}), 404
        
        # Get the API configuration
        api_config = CUSTOM_API_CONFIGS[api_id]
        base_url = api_config['baseUrl']
        path = api_config['path'] or ''  # Use empty string if path is not provided
        # --- NEW URL JOINING LOGIC ---
        clean_base_url = base_url.rstrip('/')
        clean_path = path.lstrip('/')
        if clean_path:
             api_url = f"{clean_base_url}/{clean_path}"
        else:
             api_url = clean_base_url # Use base URL directly if path is empty
        # --- END NEW LOGIC ---
        secret = api_config['secret']
        model_name = api_config['modelName']
        
        # Prepare headers based on configuration
        headers = {
            "Content-Type": "application/json"
        }
        if secret:
            headers["Authorization"] = f"Bearer {secret}"

        # Prepare payload for custom API (assuming OpenAI-like format)
        custom_payload = {
            "model": model_name,
            "messages": [
                {"role": "user", "content": prompt}
            ],
            "temperature": float(temperature),
            "stream": True
        }
        
        # 调试用，打印请求信息
        print('请求URL:', api_url)
        print('请求Headers:', headers)
        print('请求Body:', json.dumps(custom_payload, ensure_ascii=False))
        
        try:
            # Create streaming response
            def event_stream():
                try:
                    response = requests.post(
                        api_url, 
                        headers=headers, 
                        json=custom_payload, 
                        stream=True
                    )
                    response.raise_for_status()
                    
                    # Stream the response
                    for line in response.iter_lines():
                        if line:
                            decoded_line = line.decode('utf-8')
                            if decoded_line.startswith('data: '):
                                data_content = decoded_line[len('data: '):]
                                if data_content.strip() == '[DONE]':
                                    print("Stream finished.")
                                    break
                                try:
                                    chunk = json.loads(data_content)
                                    # Extract token from response format
                                    if 'choices' in chunk and len(chunk['choices']) > 0:
                                        delta = chunk['choices'][0].get('delta', {})
                                        # 兼容 ModelScope reasoning_content
                                        reasoning = delta.get('reasoning_content')
                                        token = delta.get('content')
                                        if reasoning:
                                            yield f"data: {json.dumps({'reasoning': reasoning})}\n\n"
                                        if token:
                                            yield f"data: {json.dumps({'token': token})}\n\n"
                                        
                                        finish_reason = chunk['choices'][0].get('finish_reason')
                                        if finish_reason:
                                            print(f"--> Stream finished with reason: {finish_reason}")
                                            break
                                except json.JSONDecodeError:
                                    print(f"Error decoding stream JSON chunk: {data_content}")
                            elif decoded_line.strip():
                                print(f"Received non-data line: {decoded_line}")
                except requests.exceptions.RequestException as e:
                    error_message = f"Error connecting to custom API: {e}"
                    print(error_message)
                    try:
                        if e.response is not None:
                            error_detail = e.response.json().get('error', {}).get('message', str(e))
                            error_message = f"Custom API error (Status {e.response.status_code}): {error_detail}"
                        else:
                            error_message = f"Failed to generate text via custom API. Request Error: {str(e)}"
                    except Exception:
                        error_message = f"Failed to generate text via custom API. {str(e)}"
                        if e.response is not None:
                            error_message += f" Raw Response: {e.response.text[:200]}..."
                    
                    yield f"data: {json.dumps({'error': error_message})}\n\n"
                except Exception as e:
                    error_message = f"An unexpected error occurred during custom API generation: {e}"
                    print(error_message)
                    yield f"data: {json.dumps({'error': error_message})}\n\n"
                finally:
                    print("Closing stream generator.")
            
            return Response(stream_with_context(event_stream()), mimetype='text/event-stream')
        except Exception as e:
            print(f"Error setting up custom API request: {e}")
            return jsonify({"error": f"Failed to generate text via custom API: {str(e)}"}), 500
    
    # If not a custom API, use OpenRouter (existing logic)
    if not OPENROUTER_API_KEY:
        return jsonify({"error": "请先设置OpenRouter API密钥。您可以在页面顶部的设置区域中输入并保存API密钥。"}), 401

    # --- OpenRouter API Request Logic ---
    openrouter_payload = {
        "model": model,
        "messages": [
            # OpenRouter uses a 'messages' array, typically with roles
            {"role": "user", "content": prompt}
            # Add system prompts or previous conversation history here if needed
        ],
        "temperature": float(temperature),
        "stream": True, # Enable streaming from OpenRouter
        # Add other OpenRouter parameters here if needed (e.g., max_tokens)
    }

    headers = {
        "Authorization": f"Bearer {OPENROUTER_API_KEY}",
        "Content-Type": "application/json"
        # OpenRouter might suggest adding HTTP Referer or X-Title headers
        # "HTTP-Referer": $YOUR_SITE_URL, # Optional
        # "X-Title": $YOUR_SITE_NAME # Optional
    }

    try:
        # --- Streaming Response Handling ---
        def event_stream():
            try:
                # Use the OpenRouter chat completions endpoint with streaming enabled
                response = requests.post(
                    f"{OPENROUTER_API_BASE}/chat/completions", 
                    headers=headers, 
                    json=openrouter_payload, 
                    stream=True # Enable streaming in requests
                )
                response.raise_for_status() # Raise HTTPError for bad responses (4xx or 5xx)

                # Iterate over the response lines
                for line in response.iter_lines():
                    if line:
                        decoded_line = line.decode('utf-8')
                        if decoded_line.startswith('data: '):
                            data_content = decoded_line[len('data: '):]
                            if data_content.strip() == '[DONE]': # Check for OpenRouter stream end signal
                                print("Stream finished.")
                                break
                            try:
                                chunk = json.loads(data_content)
                                # Extract token from OpenRouter's streaming format
                                # Typically: chunk['choices'][0]['delta']['content']
                                if 'choices' in chunk and len(chunk['choices']) > 0:
                                    delta = chunk['choices'][0].get('delta', {})
                                    # 兼容 ModelScope reasoning_content
                                    reasoning = delta.get('reasoning_content')
                                    token = delta.get('content')
                                    if reasoning:
                                        yield f"data: {json.dumps({'reasoning': reasoning})}\n\n"
                                    if token:
                                        yield f"data: {json.dumps({'token': token})}\n\n"
                                    
                                    finish_reason = chunk['choices'][0].get('finish_reason')
                                    if finish_reason:
                                        print(f"--> Stream finished with reason: {finish_reason}") # Enhanced logging
                                        # Optionally send a final signal
                                        # yield f"data: {json.dumps({'event': 'finish', 'reason': finish_reason})}\n\n"
                                        break # Stop sending after finish reason

                            except json.JSONDecodeError:
                                print(f"Error decoding stream JSON chunk: {data_content}")
                                # yield f"data: {json.dumps({'error': 'Invalid JSON chunk'})}\n\n"
                                # Continue to next line
                        elif decoded_line.strip(): # Log other non-empty lines for debugging
                             print(f"Received non-data line: {decoded_line}")

            except requests.exceptions.RequestException as e:
                # Handle connection errors, timeouts, etc.
                error_message = f"Error connecting to OpenRouter: {e}"
                print(error_message)
                # Try to get more specific error from response if available and it's a JSON response
                try:
                    if e.response is not None:
                        error_detail = e.response.json().get('error', {}).get('message', str(e))
                        error_message = f"OpenRouter API error (Status {e.response.status_code}): {error_detail}"
                    else:
                        error_message = f"Failed to generate text via OpenRouter. Request Error: {str(e)}"
                except Exception:
                    error_message = f"Failed to generate text via OpenRouter. {str(e)}"
                    if e.response is not None:
                        error_message += f" Raw Response: {e.response.text[:200]}..."
                
                # Send error back via SSE
                yield f"data: {json.dumps({'error': error_message})}\n\n"
            except Exception as e:
                error_message = f"An unexpected error occurred during OpenRouter generation: {e}"
                print(error_message)
                yield f"data: {json.dumps({'error': error_message})}\n\n"
            finally:
                # Ensure a final signal is sent if needed, or just close
                print("Closing stream generator.")
                # yield f"data: [DONE]\n\n" # Optionally send a custom end signal

        # Return the streaming response
        return Response(stream_with_context(event_stream()), mimetype='text/event-stream')

    except requests.exceptions.RequestException as e:
        # Handle connection errors, timeouts, etc.
        error_message = f"Error connecting to OpenRouter: {e}"
        print(error_message)
            # Try to get more specific error from response if available and it's a JSON response
        # --- This error handling is for the initial request, not the stream itself ---
        # --- Stream errors are handled within event_stream --- 
        try:
            if e.response is not None:
                error_detail = e.response.json().get('error', {}).get('message', str(e))
                error_message = f"OpenRouter API error (Status {e.response.status_code}): {error_detail}"
            else:
                error_message = f"Failed to generate text via OpenRouter. Request Error: {str(e)}"
        except Exception:
            error_message = f"Failed to generate text via OpenRouter. {str(e)}"
            if e.response is not None:
                error_message += f" Raw Response: {e.response.text[:200]}..."

        return jsonify({"error": error_message}), 500
    except json.JSONDecodeError:
            print(f"Error decoding JSON response from OpenRouter generation: {response.text}")
            return jsonify({"error": "Invalid response format from OpenRouter generation."}), 500
    except Exception as e:
        print(f"An unexpected error occurred during OpenRouter generation: {e}")
        return jsonify({"error": f"An unexpected error occurred: {str(e)}"}), 500

# --- Export Endpoints (Server-Side) ---

def generate_export_filename(title, extension):
    """Generates a sanitized filename with title and timestamp."""
    # Sanitize title (allow alphanumeric and CJK)
    sanitized_title = ''.join(c for c in title if c.isalnum() or '\u4e00' <= c <= '\u9fff')
    sanitized_title = sanitized_title or "Untitled" # Fallback if title is empty after sanitizing
    timestamp = datetime.datetime.now().strftime("%Y-%m-%dT%H-%M-%S")
    return f"{sanitized_title}_{timestamp}.{extension}"

@app.route('/api/novels/<novel_id>/export/<file_format>', methods=['POST'])
def export_novel_serverside(novel_id, file_format):
    """API endpoint to export a novel to a file on the server."""
    if novel_id not in novels_db:
        return jsonify({"error": "Novel not found"}), 404

    novel_data = novels_db[novel_id]
    title = novel_data.get('title', 'Untitled')
    content = novel_data.get('content', '')
    characters = novel_data.get('characters', '')
    knowledge = novel_data.get('knowledge', '')
    style_prompt = novel_data.get('style_prompt', '')
    character_tags = novel_data.get('character_tags', [])
    glossary_tags = novel_data.get('glossary_tags', [])
    style_tags = novel_data.get('style_tags', [])

    try:
        if file_format == 'txt':
            filename = generate_export_filename(title, 'txt')
            filepath = os.path.join(EXPORT_DIR, filename)
            txt_content = f"标题: {title}\n\n"
            if characters: txt_content += f"角色设定:\n{characters}\n\n"
            if style_prompt: txt_content += f"风格提示:\n{style_prompt}\n\n"
            if knowledge: txt_content += f"关联知识:\n{knowledge}\n\n"
            txt_content += f"正文内容:\n{content}"
            with open(filepath, 'w', encoding='utf-8') as f:
                f.write(txt_content)

        elif file_format == 'json':
            filename = generate_export_filename(title, 'json')
            filepath = os.path.join(EXPORT_DIR, filename)
            export_data = {
                'title': title,
                'content': content,
                'characters': characters,
                'knowledge': knowledge,
                'style_prompt': style_prompt,
                'character_tags': character_tags,
                'glossary_tags': glossary_tags,
                'style_tags': style_tags
            }
            with open(filepath, 'w', encoding='utf-8') as f:
                json.dump(export_data, f, ensure_ascii=False, indent=4)

        elif file_format == 'docx':
            filename = generate_export_filename(title, 'docx')
            filepath = os.path.join(EXPORT_DIR, filename)

            document = Document()
            document.add_heading(title, level=0) # Title
            document.add_paragraph()

            if characters:
                document.add_heading('角色设定', level=2)
                for line in characters.split('\n'):
                    document.add_paragraph(line)
                document.add_paragraph()
            if style_prompt:
                document.add_heading('风格提示', level=2)
                for line in style_prompt.split('\n'):
                    document.add_paragraph(line)
                document.add_paragraph()
            if knowledge:
                document.add_heading('关联知识', level=2)
                for line in knowledge.split('\n'):
                    document.add_paragraph(line)
                document.add_paragraph()

            document.add_heading('正文内容', level=1)
            for line in content.split('\n'): # Add content line by line
                document.add_paragraph(line)

            document.save(filepath)

        else:
            return jsonify({"error": "Invalid file format requested"}), 400

        print(f"Exported novel '{title}' ({novel_id}) to {filepath}")
        # Return path relative to server script for confirmation
        try:
            relative_path = os.path.relpath(filepath, os.path.dirname(__file__))
        except ValueError:
            relative_path = filepath
        return jsonify({"message": f"文件已成功导出至服务器: {relative_path}", "filepath": relative_path})

    except Exception as e:
        print(f"Error exporting novel {novel_id} to {file_format}: {e}")
        return jsonify({"error": f"导出文件时发生错误: {str(e)}"}), 500

@app.route('/api/novels/<novel_id>/export-notes', methods=['POST'])
def export_novel_notes(novel_id):
    """API endpoint to export novel notes to a file in the exports folder."""
    data = request.get_json()
    if not data or 'content' not in data:
        logging.error(f"[笔记导出] 未提供content，novel_id={novel_id}")
        return jsonify({"error": "No content provided"}), 400

    try:
        # Create exports directory if it doesn't exist
        os.makedirs(EXPORT_DIR, exist_ok=True)

        # Generate filename with timestamp
        from datetime import datetime
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f'novel_notes_{novel_id}_{timestamp}.txt'
        filepath = os.path.join(EXPORT_DIR, filename)

        # Write content to file
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(data['content'])

        logging.info(f"[笔记导出] 成功: {filepath}")
        return jsonify({
            "message": f"笔记已成功导出到: {filename}",
            "filepath": filepath
        })

    except Exception as e:
        logging.error(f"[笔记导出] Error exporting notes for novel {novel_id}: {str(e)}")
        return jsonify({"error": f"导出笔记失败: {str(e)}"}), 500

# --- Import Endpoint ---

@app.route('/api/novels/import', methods=['POST'])
def import_novel():
    """API endpoint to import a novel from uploaded file content."""
    data = request.get_json()
    if not data or 'filename' not in data or 'content' not in data:
        return jsonify({"error": "Missing filename or content"}), 400

    filename = data['filename']
    file_content = data['content']
    file_ext = filename.split('.')[-1].lower()

    novel_data = {
        "title": "Untitled Import",
        "content": "",
        "characters": "",
        "knowledge": "",
        "style_prompt": "",
        "character_tags": [],
        "glossary_tags": [],
        "style_tags": []
    }

    try:
        if file_ext == 'json':
            # For JSON files, check if the content is already a JSON string or needs parsing
            try:
                # Try parsing the JSON content
                imported_data = json.loads(strip_code_block(file_content))
                novel_data['title'] = imported_data.get('title', f"Imported {filename}")
                novel_data['content'] = imported_data.get('content', '')
                novel_data['characters'] = imported_data.get('characters', '')
                novel_data['knowledge'] = imported_data.get('knowledge', '')
                novel_data['style_prompt'] = imported_data.get('style_prompt', '')
                # Import tag arrays if present
                novel_data['character_tags'] = imported_data.get('character_tags', [])
                novel_data['glossary_tags'] = imported_data.get('glossary_tags', [])
                novel_data['style_tags'] = imported_data.get('style_tags', [])
            except json.JSONDecodeError as json_error:
                # More detailed error message for JSON parsing issues
                detail = str(json_error)
                print(f"JSON parse error: {detail}. Content preview: {file_content[:100]}...")
                return jsonify({"error": f"无效的JSON格式: {detail}"}), 400

        elif file_ext == 'txt':
            lines = file_content.splitlines()
            current_section = None
            content_buffer = []
            characters_buffer = []
            knowledge_buffer = []
            style_buffer = []

            for line in lines:
                line_strip = line.strip()
                if line_strip.startswith("标题:"):
                    novel_data['title'] = line_strip.split(":", 1)[1].strip() or f"Imported {filename}"
                    current_section = None
                elif line_strip == "角色设定:":
                    current_section = "characters"
                elif line_strip == "风格提示:":
                    current_section = "style_prompt"
                elif line_strip == "关联知识:":
                    current_section = "knowledge"
                elif line_strip == "正文内容:":
                    current_section = "content"
                elif current_section == "characters":
                    characters_buffer.append(line)
                elif current_section == "style_prompt":
                    style_buffer.append(line)
                elif current_section == "knowledge":
                    knowledge_buffer.append(line)
                elif current_section == "content":
                    content_buffer.append(line)
                elif current_section is None and novel_data['title'] == "Untitled Import": # Try to find title if not explicitly tagged
                     if line_strip: # Use first non-empty line as title if no header found yet
                         novel_data['title'] = line_strip
                         current_section = "maybe_content" # Assume following lines might be content

                elif current_section == "maybe_content": # If title was inferred, assume rest is content until a proper header
                     content_buffer.append(line)


            novel_data['content'] = "\n".join(content_buffer).strip()
            novel_data['characters'] = "\n".join(characters_buffer).strip()
            novel_data['knowledge'] = "\n".join(knowledge_buffer).strip()
            novel_data['style_prompt'] = "\n".join(style_buffer).strip()

            # Basic fallback if no headers found at all
            if not novel_data['content'] and not novel_data['characters'] and not novel_data['knowledge'] and not novel_data['style_prompt']:
                if novel_data['title'] == "Untitled Import":
                    novel_data['title'] = f"Imported {filename}"
                novel_data['content'] = file_content # Put everything in content
                
        elif file_ext == 'docx':
            # For Word documents, we need to convert base64 content to binary
            import base64
            
            try:
                # Convert base64 content to binary
                binary_content = base64.b64decode(file_content)
                
                # Read the Word document
                doc = Document(io.BytesIO(binary_content))
                
                paragraphs = [p.text for p in doc.paragraphs]
                
                # Try to identify sections based on headings
                current_section = "maybe_content"
                content_buffer = []
                characters_buffer = []
                knowledge_buffer = []
                style_buffer = []
                
                # First check if the document has a title (first paragraph or heading)
                if paragraphs and paragraphs[0].strip():
                    novel_data["title"] = paragraphs[0].strip()
                    # Skip the title in further processing
                    paragraphs = paragraphs[1:]
                else:
                    novel_data["title"] = f"Imported {filename}"
                
                for p in paragraphs:
                    p_strip = p.strip()
                    
                    # Check for section headers
                    if p_strip.lower() in ["角色设定", "角色"]:
                        current_section = "characters"
                        continue
                    elif p_strip.lower() in ["风格提示", "风格"]:
                        current_section = "style_prompt"
                        continue
                    elif p_strip.lower() in ["关联知识", "知识"]:
                        current_section = "knowledge"
                        continue
                    elif p_strip.lower() in ["正文内容", "正文", "内容"]:
                        current_section = "content"
                        continue
                    
                    # Add paragraph to appropriate buffer
                    if current_section == "characters":
                        characters_buffer.append(p)
                    elif current_section == "style_prompt":
                        style_buffer.append(p)
                    elif current_section == "knowledge":
                        knowledge_buffer.append(p)
                    elif current_section == "content" or current_section == "maybe_content":
                        content_buffer.append(p)
                
                # Join the buffers and store in novel_data
                novel_data['content'] = "\n".join(content_buffer).strip()
                novel_data['characters'] = "\n".join(characters_buffer).strip()
                novel_data['knowledge'] = "\n".join(knowledge_buffer).strip()
                novel_data['style_prompt'] = "\n".join(style_buffer).strip()
                
                # If no structured content was found, put everything in content
                if not novel_data['content'] and not novel_data['characters'] and not novel_data['knowledge'] and not novel_data['style_prompt']:
                    novel_data['content'] = "\n".join(paragraphs).strip()
            except base64.binascii.Error as base64_error:
                print(f"Base64 decode error: {str(base64_error)}. Content preview: {file_content[:100]}...")
                return jsonify({"error": f"无法解码Word文档: {str(base64_error)}"}), 400
            except Exception as docx_error:
                print(f"DOCX processing error: {str(docx_error)}")
                return jsonify({"error": f"处理Word文档时出错: {str(docx_error)}"}), 400

        else:
            return jsonify({"error": "不支持的文件格式 (仅支持 .txt, .json, .docx)"}), 400

        # 统一设置标题为"用户导入"
        novel_data['title'] = "用户导入"

        # Create new novel entry
        novel_id = str(uuid.uuid4())
        novels_db[novel_id] = novel_data
        print(f"Imported novel: {novel_id} - {novel_data['title']}")
        save_novels()  # 保存小说数据

        # Return the details of the imported novel
        return jsonify({
            "message": "Novel imported successfully",
            "id": novel_id,
            "title": novel_data["title"],
            "content": novel_data["content"],
            "characters": novel_data["characters"],
            "knowledge": novel_data["knowledge"],
            "style_prompt": novel_data["style_prompt"]
        }), 201

    except json.JSONDecodeError as json_error:
        error_message = f"无效的JSON格式: {str(json_error)}"
        print(f"JSON parsing error: {error_message}")
        return jsonify({"error": error_message}), 400
    except Exception as e:
        error_message = f"导入小说时发生错误: {str(e)}"
        print(f"Error importing novel from {filename}: {e}")
        return jsonify({"error": error_message}), 500

# --- Character Library Endpoints ---

@app.route('/api/characters', methods=['GET'])
def get_characters():
    """API endpoint to list all characters in the library."""
    # Return a list of characters with id and name
    character_list = [{"id": char_id, "name": data.get("name", "未命名"), "description": data.get("description", "")[:50] + "..." if len(data.get("description", "")) > 50 else data.get("description", "")} 
                     for char_id, data in characters_db.items()]
    return jsonify(character_list)

@app.route('/api/characters', methods=['POST'])
def create_character():
    """API endpoint to create a new character."""
    data = request.get_json()
    if not data or 'name' not in data:
        return jsonify({"error": "角色名称是必填项"}), 400

    character_id = str(uuid.uuid4()) # Generate a unique ID
    new_character = {
        "name": data['name'],
        "description": data.get('description', ''), # Character description
        "details": data.get('details', {})  # Additional details (could be skills, backstory, etc.)
    }
    characters_db[character_id] = new_character
    print(f"Created character: {character_id} - {data['name']}")
    save_characters()  # 保存角色数据
    
    return jsonify({
        "id": character_id, 
        "name": new_character["name"],
        "description": new_character["description"],
        "details": new_character["details"]
    }), 201 # 201 Created status

@app.route('/api/characters/<character_id>', methods=['GET'])
def get_character(character_id):
    """API endpoint to get details of a specific character."""
    character = characters_db.get(character_id)
    if character:
        return jsonify({
            "id": character_id,
            "name": character.get("name", "未命名"),
            "description": character.get("description", ""),
            "details": character.get("details", {})
        })
    else:
        return jsonify({"error": "角色不存在"}), 404

@app.route('/api/characters/<character_id>', methods=['PUT', 'PATCH'])
def update_character(character_id):
    """API endpoint to update a character's information."""
    if character_id not in characters_db:
        return jsonify({"error": "角色不存在"}), 404
        
    data = request.get_json()
    if not data:
        return jsonify({"error": "未提供更新数据"}), 400
        
    # Update fields provided in the request
    if 'name' in data:
        characters_db[character_id]['name'] = data['name']
    if 'description' in data:
        characters_db[character_id]['description'] = data['description']
    if 'details' in data:
        characters_db[character_id]['details'] = data['details']
        
    save_characters()  # 保存角色数据
    return jsonify({
        "message": "角色已更新", 
        "id": character_id
    })

@app.route('/api/characters/<character_id>', methods=['DELETE'])
def delete_character(character_id):
    """API endpoint to delete a specific character."""
    if character_id in characters_db:
        deleted_name = characters_db[character_id].get('name', '未命名')
        del characters_db[character_id]
        print(f"Deleted character: {character_id} - {deleted_name}")
        save_characters()  # 保存角色数据
        return jsonify({"message": f"角色 '{deleted_name}' 已成功删除"})
    else:
        return jsonify({"error": "角色不存在"}), 404

# --- Glossary Endpoints ---

@app.route('/api/glossary', methods=['GET'])
def get_glossary_entries():
    """API endpoint to list all glossary entries."""
    # Return a list of glossary entries with id, term and brief description
    glossary_list = [{"id": entry_id, 
                     "term": data.get("term", "未命名词条"), 
                     "category": data.get("category", ""),
                     "description": data.get("description", "")[:50] + "..." if len(data.get("description", "")) > 50 else data.get("description", "")} 
                    for entry_id, data in glossary_db.items()]
    return jsonify(glossary_list)

@app.route('/api/glossary', methods=['POST'])
def create_glossary_entry():
    """API endpoint to create a new glossary entry."""
    data = request.get_json()
    if not data or 'term' not in data:
        return jsonify({"error": "词条名称是必填项"}), 400

    entry_id = str(uuid.uuid4()) # Generate a unique ID
    new_entry = {
        "term": data['term'],
        "category": data.get('category', ''), # Category for organization
        "description": data.get('description', ''), # Description of the term
        "details": data.get('details', {})  # Additional details if needed
    }
    glossary_db[entry_id] = new_entry
    print(f"Created glossary entry: {entry_id} - {data['term']}")
    save_glossary()  # 保存词条数据
    
    return jsonify({
        "id": entry_id, 
        "term": new_entry["term"],
        "category": new_entry["category"],
        "description": new_entry["description"],
        "details": new_entry["details"]
    }), 201 # 201 Created status

@app.route('/api/glossary/<entry_id>', methods=['GET'])
def get_glossary_entry(entry_id):
    """API endpoint to get details of a specific glossary entry."""
    entry = glossary_db.get(entry_id)
    if entry:
        return jsonify({
            "id": entry_id,
            "term": entry.get("term", "未命名词条"),
            "category": entry.get("category", ""),
            "description": entry.get("description", ""),
            "details": entry.get("details", {})
        })
    else:
        return jsonify({"error": "词条不存在"}), 404

@app.route('/api/glossary/<entry_id>', methods=['PUT', 'PATCH'])
def update_glossary_entry(entry_id):
    """API endpoint to update a glossary entry."""
    if entry_id not in glossary_db:
        return jsonify({"error": "词条不存在"}), 404

    data = request.get_json()
    if not data:
        return jsonify({"error": "未提供更新数据"}), 400

    # Update fields provided in the request
    if 'term' in data:
        glossary_db[entry_id]['term'] = data['term']
    if 'category' in data:
        glossary_db[entry_id]['category'] = data['category']
    if 'description' in data:
        glossary_db[entry_id]['description'] = data['description']
    if 'details' in data:
        glossary_db[entry_id]['details'] = data['details']

    save_glossary()  # 保存词条数据
    return jsonify({
        "message": "词条已更新", 
        "id": entry_id
    })

@app.route('/api/glossary/<entry_id>', methods=['DELETE'])
def delete_glossary_entry(entry_id):
    """API endpoint to delete a specific glossary entry."""
    if entry_id in glossary_db:
        deleted_term = glossary_db[entry_id].get('term', '未命名词条')
        del glossary_db[entry_id]
        print(f"Deleted glossary entry: {entry_id} - {deleted_term}")
        save_glossary()  # 保存词条数据
        return jsonify({"message": f"词条 '{deleted_term}' 已成功删除"})
    else:
        return jsonify({"error": "词条不存在"}), 404

@app.route('/api/novels/<novel_id>/glossary', methods=['GET'])
def get_novel_glossary(novel_id):
    """API endpoint to get glossary entries associated with a specific novel."""
    if novel_id not in novels_db:
        return jsonify({"error": "小说不存在"}), 404
        
    # In a more advanced implementation, we might have a relationship table
    # For now, we'll return all glossary entries (could be filtered by novel later)
    return get_glossary_entries()

# --- Style Library Endpoints ---

@app.route('/api/styles', methods=['GET'])
def get_styles():
    """API endpoint to list all style entries."""
    # Return a list of style entries with id, name and brief description
    style_list = [{"id": style_id, 
                  "name": data.get("name", "未命名风格"), 
                  "category": data.get("category", ""),
                  "content": data.get("content", "")[:50] + "..." if len(data.get("content", "")) > 50 else data.get("content", "")} 
                 for style_id, data in styles_db.items()]
    return jsonify(style_list)

@app.route('/api/styles', methods=['POST'])
def create_style():
    """API endpoint to create a new style entry."""
    data = request.get_json()
    if not data or 'name' not in data:
        return jsonify({"error": "风格名称是必填项"}), 400

    style_id = str(uuid.uuid4()) # Generate a unique ID
    new_style = {
        "name": data['name'],
        "category": data.get('category', ''), # Category for organization
        "content": data.get('content', ''), # The actual style prompt content
        "details": data.get('details', {})  # Additional details if needed
    }
    styles_db[style_id] = new_style
    print(f"Created style entry: {style_id} - {data['name']}")
    save_styles()  # 保存风格数据
    
    return jsonify({
        "id": style_id, 
        "name": new_style["name"],
        "category": new_style["category"],
        "content": new_style["content"],
        "details": new_style["details"]
    }), 201 # 201 Created status

@app.route('/api/styles/import-md', methods=['POST'])
def import_md_as_style():
    """API endpoint to import a markdown file from the prompt folder as a style entry."""
    data = request.get_json()
    if not data or 'filename' not in data:
        return jsonify({"error": "文件名是必填项"}), 400

    try:
        # Determine the prompt folder path
        if getattr(sys, 'frozen', False):
            prompt_folder = os.path.join(os.path.dirname(sys.executable), 'prompt')
        else:
            prompt_folder = os.path.join(os.path.dirname(__file__), 'prompt')
        
        # Ensure the filename has the .md extension
        filename = data['filename']
        if not filename.lower().endswith('.md'):
            filename += '.md'
        
        file_path = os.path.join(prompt_folder, filename)
        
        if not os.path.exists(file_path):
            return jsonify({"error": f"找不到文件: {filename}"}), 404
        
        # Read the markdown file content
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # Extract name from filename (remove extension)
        name = os.path.splitext(filename)[0]
        
        # Create a new style entry
        style_id = str(uuid.uuid4())
        new_style = {
            "name": name,
            "category": data.get('category', '导入提示词'),  # Default category or user-provided
            "content": content,
            "details": {
                "imported": True,
                "source": file_path,
                "import_date": datetime.datetime.now().isoformat()
            }
        }
        
        styles_db[style_id] = new_style
        print(f"Imported style from MD file: {style_id} - {name}")
        save_styles()  # 保存风格数据
        
        return jsonify({
            "id": style_id, 
            "name": new_style["name"],
            "category": new_style["category"],
            "content": new_style["content"],
            "details": new_style["details"]
        }), 201
    except Exception as e:
        return jsonify({"error": f"导入提示词文件时出错: {str(e)}"}), 500

@app.route('/api/styles/list-md-files', methods=['GET'])
def list_md_files():
    """API endpoint to list all markdown files in the prompt folder."""
    try:
        # Determine the prompt folder path
        if getattr(sys, 'frozen', False):
            prompt_folder = os.path.join(os.path.dirname(sys.executable), 'prompt')
        else:
            prompt_folder = os.path.join(os.path.dirname(__file__), 'prompt')
        
        # Get all .md files in the prompt folder
        md_files = [f for f in os.listdir(prompt_folder) if f.lower().endswith('.md')]
        
        return jsonify(md_files)
    except Exception as e:
        return jsonify({"error": f"列出提示词文件时出错: {str(e)}"}), 500

@app.route('/api/styles/<style_id>', methods=['GET'])
def get_style(style_id):
    """API endpoint to get details of a specific style entry."""
    style = styles_db.get(style_id)
    if style:
        return jsonify({
            "id": style_id,
            "name": style.get("name", "未命名风格"),
            "category": style.get("category", ""),
            "content": style.get("content", ""),
            "details": style.get("details", {})
        })
    else:
        return jsonify({"error": "风格不存在"}), 404

@app.route('/api/styles/<style_id>', methods=['PUT', 'PATCH'])
def update_style(style_id):
    """API endpoint to update a style entry."""
    if style_id not in styles_db:
        return jsonify({"error": "风格不存在"}), 404

    data = request.get_json()
    if not data:
        return jsonify({"error": "未提供更新数据"}), 400

    # Update fields provided in the request
    if 'name' in data:
        styles_db[style_id]['name'] = data['name']
    if 'category' in data:
        styles_db[style_id]['category'] = data['category']
    if 'content' in data:
        styles_db[style_id]['content'] = data['content']
    if 'details' in data:
        styles_db[style_id]['details'] = data['details']

    save_styles()  # 保存风格数据
    return jsonify({
        "message": "风格已更新", 
        "id": style_id
    })

@app.route('/api/styles/<style_id>', methods=['DELETE'])
def delete_style(style_id):
    """API endpoint to delete a specific style entry."""
    if style_id in styles_db:
        deleted_name = styles_db[style_id].get('name', '未命名风格')
        del styles_db[style_id]
        print(f"Deleted style entry: {style_id} - {deleted_name}")
        save_styles()  # 保存风格数据
        return jsonify({"message": f"风格 '{deleted_name}' 已成功删除"})
    else:
        return jsonify({"error": "风格不存在"}), 404

@app.route('/api/novels/<novel_id>/styles', methods=['GET'])
def get_novel_styles(novel_id):
    """API endpoint to get style entries associated with a specific novel."""
    if novel_id not in novels_db:
        return jsonify({"error": "小说不存在"}), 404
        
    # In a more advanced implementation, we might have a relationship table
    # For now, we'll return all style entries (could be filtered by novel later)
    return get_styles()

@app.route('/api/games', methods=['GET'])
def get_games():
    """API endpoint to list all games."""
    game_list = [{"id": game_id, "title": data.get("title", "Untitled Game")} for game_id, data in games_db.items()]
    return jsonify(game_list)

@app.route('/api/games', methods=['POST'])
def create_game():
    """API endpoint to create a new game."""
    data = request.get_json()
    
    game_id = str(uuid.uuid4()) # Generate a unique ID
    new_game = {
        "title": data.get('title', 'New Game'),
        "character": data.get('character', ''),
        "world": data.get('world', ''),
        "style": data.get('style', ''),
        "chat_history": data.get('chat_history', []),
        "created_at": datetime.datetime.now().isoformat()
    }
    games_db[game_id] = new_game
    print(f"Created game: {game_id} - {new_game['title']}")
    save_games()  # 保存游戏数据
    return jsonify({"id": game_id, "title": new_game["title"]}), 201 # 201 Created status

@app.route('/api/games/<game_id>', methods=['GET'])
def get_game(game_id):
    """API endpoint to get details of a specific game."""
    game = games_db.get(game_id)
    if game:
        return jsonify({
            "id": game_id,
            "title": game.get("title", "Untitled"),
            "character": game.get("character", ""),
            "world": game.get("world", ""),
            "style": game.get("style", ""),
            "chat_history": game.get("chat_history", []),
            "created_at": game.get("created_at", "")
        })
    else:
        return jsonify({"message": "Game not found"}), 404

@app.route('/api/games/<game_id>', methods=['PUT', 'PATCH'])
def update_game(game_id):
    """API endpoint to update a game's settings and history."""
    if game_id not in games_db:
        return jsonify({"message": "Game not found"}), 404

    data = request.get_json()
    if not data:
        return jsonify({"message": "No update data provided"}), 400

    # Update fields provided in the request
    if 'title' in data:
        games_db[game_id]['title'] = data['title']
    if 'character' in data:
        games_db[game_id]['character'] = data['character']
    if 'world' in data:
        games_db[game_id]['world'] = data['world']
    if 'style' in data:
        games_db[game_id]['style'] = data['style']
    if 'chat_history' in data:
        games_db[game_id]['chat_history'] = data['chat_history']
    
    save_games()  # 保存游戏数据
    return jsonify({"message": "Game updated successfully"})

@app.route('/api/games/<game_id>', methods=['DELETE'])
def delete_game(game_id):
    """API endpoint to delete a game."""
    if game_id not in games_db:
        return jsonify({"message": "Game not found"}), 404
    
    deleted_title = games_db[game_id].get('title', 'Unknown')
    del games_db[game_id]
    save_games()  # 保存游戏数据
    
    return jsonify({"message": f"Game '{deleted_title}' deleted successfully"})

# --- Export Endpoints (Chat History) ---

@app.route('/api/chat/export', methods=['POST'])
def export_chat_history():
    """API endpoint to export chat history to a JSON file."""
    data = request.get_json()
    if not data or 'chat_history' not in data or 'settings' not in data:
        return jsonify({"error": "Missing required chat data"}), 400

    chat_history = data['chat_history']
    settings = data['settings']
    
    try:
        # Generate a filename with timestamp
        timestamp = datetime.datetime.now().strftime("%Y-%m-%dT%H-%M-%S")
        filename = f"Chat_Export_{timestamp}.json"
        filepath = os.path.join(EXPORT_DIR, filename)
        
        # Prepare export data
        export_data = {
            'timestamp': timestamp,
            'settings': settings,
            'chat_history': chat_history
        }
        
        # Write to file
        with open(filepath, 'w', encoding='utf-8') as f:
            json.dump(export_data, f, ensure_ascii=False, indent=4)
        
        print(f"Exported chat history to {filepath}")
        # Return path relative to server script for confirmation
        try:
            relative_path = os.path.relpath(filepath, os.path.dirname(__file__))
        except ValueError:
            relative_path = filepath
        return jsonify({
            "message": f"聊天记录已成功导出至服务器: {relative_path}", 
            "filepath": relative_path
        })
        
    except Exception as e:
        print(f"Error exporting chat history: {e}")
        return jsonify({"error": f"导出聊天记录时发生错误: {str(e)}"}), 500

@app.route('/api/chat/import', methods=['POST'])
def import_chat_history():
    """API endpoint to import chat history from a JSON file content."""
    data = request.get_json()
    if not data or 'file_content' not in data:
        return jsonify({"error": "Missing file content"}), 400

    file_content = data['file_content']
    
    try:
        # Parse the JSON content
        imported_data = json.loads(strip_code_block(file_content))
        
        # Validate the imported data
        if 'chat_history' not in imported_data or 'settings' not in imported_data:
            return jsonify({"error": "无效的聊天记录文件格式"}), 400
        
        return jsonify({
            "message": "聊天记录已成功导入",
            "chat_data": imported_data
        })
        
    except json.JSONDecodeError as json_error:
        # More detailed error message for JSON parsing issues
        detail = str(json_error)
        print(f"JSON parse error: {detail}. Content preview: {file_content[:100]}...")
        return jsonify({"error": f"无效的JSON格式: {detail}"}), 400
    except Exception as e:
        print(f"Error importing chat history: {e}")
        return jsonify({"error": f"导入聊天记录时发生错误: {str(e)}"}), 500

@app.route('/api/custom-apis', methods=['GET'])
def get_custom_apis():
    """API endpoint to get the list of saved custom APIs."""
    try:
        return jsonify(list(CUSTOM_API_CONFIGS.values()))
    except Exception as e:
        print(f"Error fetching custom APIs: {e}")
        return jsonify({"error": "Could not retrieve custom API list."}), 500

@app.route('/api/custom-apis', methods=['POST'])
def add_custom_api():
    """API endpoint to add a new custom API configuration."""
    data = request.get_json()
    if not data or 'name' not in data or 'baseUrl' not in data or 'modelName' not in data:
        return jsonify({"error": "Missing required fields"}), 400
    
    api_id = str(uuid.uuid4())
    new_api = {
        "id": api_id,
        "name": data['name'],
        "baseUrl": data['baseUrl'],
        "path": data.get('path', ''),
        "secret": data.get('secret', ''),
        "modelName": data['modelName']
    }
    
    CUSTOM_API_CONFIGS[api_id] = new_api
    print(f"Added custom API: {api_id} - {data['name']} - {data['baseUrl']}")
    
    # 保存到文件
    save_api_configs()
    
    return jsonify(new_api), 201

@app.route('/api/custom-apis/<api_id>', methods=['PUT'])
def update_custom_api(api_id):
    """API endpoint to update an existing custom API configuration."""
    if api_id not in CUSTOM_API_CONFIGS:
        return jsonify({"error": "API configuration not found"}), 404
    
    data = request.get_json()
    if not data:
        return jsonify({"error": "No update data provided"}), 400
    
    # Update fields
    CUSTOM_API_CONFIGS[api_id].update({
        "name": data.get('name', CUSTOM_API_CONFIGS[api_id]['name']),
        "baseUrl": data.get('baseUrl', CUSTOM_API_CONFIGS[api_id]['baseUrl']),
        "path": data.get('path', CUSTOM_API_CONFIGS[api_id]['path']),
        "secret": data.get('secret', CUSTOM_API_CONFIGS[api_id]['secret']),
        "modelName": data.get('modelName', CUSTOM_API_CONFIGS[api_id]['modelName'])
    })
    
    print(f"Updated custom API: {api_id}")
    
    # 保存到文件
    save_api_configs()
    
    return jsonify(CUSTOM_API_CONFIGS[api_id])

@app.route('/api/custom-apis/<api_id>', methods=['DELETE'])
def delete_custom_api(api_id):
    """API endpoint to delete a custom API configuration."""
    if api_id not in CUSTOM_API_CONFIGS:
        return jsonify({"error": "API configuration not found"}), 404
    
    deleted_name = CUSTOM_API_CONFIGS[api_id]['name']
    del CUSTOM_API_CONFIGS[api_id]
    print(f"Deleted custom API: {api_id} - {deleted_name}")
    
    # 保存到文件
    save_api_configs()
    
    return jsonify({"message": f"Custom API '{deleted_name}' deleted successfully"})

@app.route('/api/novel-types', methods=['GET'])
def get_novel_types():
    """API endpoint to get novel types and subtypes from the 小说类型.txt file."""
    try:
        novel_types = {}
        
        # Read the 小说类型.txt file
        novel_types_file = os.path.join(DATA_DIR, '小说类型.txt')
        if not os.path.exists(novel_types_file):
            logging.error(f"小说类型文件不存在: {novel_types_file}")
            return jsonify({"error": f"小说类型文件不存在: {novel_types_file}"}), 500
        with open(novel_types_file, 'r', encoding='utf-8') as file:
            lines = file.readlines()
        
        # Parse each line
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            # Extract main type and subtypes
            if '类型选项及展开子选项：' in line:
                parts = line.split('类型选项及展开子选项：')
                main_type = parts[0].strip()
                
                # Extract subtypes
                subtypes_text = parts[1].strip()
                if subtypes_text.endswith('。'):
                    subtypes_text = subtypes_text[:-1]  # Remove trailing period
                
                # Split by commas or Chinese commas
                subtypes = []
                for subtype in re.split('、|,|，', subtypes_text):
                    clean_subtype = subtype.strip()
                    if clean_subtype and clean_subtype not in ('包括', '等'):
                        subtypes.append(clean_subtype)
                
                novel_types[main_type] = subtypes
        
        return jsonify(novel_types)
    except Exception as e:
        print(f"Error loading novel types: {str(e)}")
        return jsonify({"error": f"Failed to load novel types: {str(e)}"}), 500

@app.route('/api/save-text-file', methods=['POST'])
def save_text_file():
    """API endpoint to save text file to the exports directory."""
    data = request.get_json()
    if not data or 'content' not in data or 'filename' not in data:
        return jsonify({"error": "Missing required fields: content or filename"}), 400
    
    content = data['content']
    filename = data['filename']
    
    # Sanitize filename (allow alphanumeric, CJK characters, and some symbols)
    sanitized_filename = ''.join(c for c in filename if c.isalnum() or '\u4e00' <= c <= '\u9fff' or c in '-_. ')
    sanitized_filename = sanitized_filename or "unnamed_file.txt" # Fallback if filename is empty after sanitizing
    
    try:
        # Ensure the exports directory exists
        if not os.path.exists(EXPORT_DIR):
            os.makedirs(EXPORT_DIR)
        
        # Create the full path
        filepath = os.path.join(EXPORT_DIR, sanitized_filename)
        
        # Write the content to the file
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(content)
        
        # Return success response with relative path
        try:
            relative_path = os.path.relpath(filepath, os.path.dirname(__file__))
        except ValueError:
            relative_path = filepath
        return jsonify({
            "message": "文件已成功保存",
            "filepath": relative_path
        })
        
    except Exception as e:
        error_message = f"保存文件时出错: {str(e)}"
        print(error_message)
        return jsonify({"error": error_message}), 500

@app.route('/api/prompt-editor/export-template', methods=['POST'])
def export_prompt_template():
    """API endpoint to save prompt editor content as a .md template file in the model directory."""
    data = request.get_json()
    if not data or 'content' not in data or 'filename' not in data:
        return jsonify({"error": "Missing required fields: content or filename"}), 400

    content = data['content']
    filename = data['filename']

    # Sanitize filename (similar to other export functions)
    # Allow alphanumeric, CJK characters, and some symbols. Ensure .md extension.
    sanitized_filename = ''.join(c for c in filename if c.isalnum() or '\\u4e00' <= c <= '\\u9fff' or c in '-_. ')
    if not sanitized_filename: # Fallback if filename is empty after sanitizing
        sanitized_filename = "prompt_template"
    
    if not sanitized_filename.lower().endswith('.md'):
        sanitized_filename += '.md'
    
    try:
        # Ensure the model directory exists (it should be created on app start, but double-check)
        if not os.path.exists(MODEL_DIR):
            os.makedirs(MODEL_DIR)
        
        filepath = os.path.join(MODEL_DIR, sanitized_filename)
        
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(content)
        
        print(f"Exported prompt template to {filepath}")
        
        # Return path relative to server script for confirmation (optional, but consistent)
        try:
            relative_path = os.path.relpath(filepath, os.path.dirname(__file__))
        except ValueError: # Handles cases where paths are on different drives (e.g., during PyInstaller build in temp)
            relative_path = filepath 
            
        return jsonify({
            "message": f"模板已成功导出至服务器 'model' 文件夹: {sanitized_filename}",
            "filepath": relative_path # Sending back the relative path for consistency
        })
        
    except Exception as e:
        error_message = f"导出模板文件时出错: {str(e)}"
        print(error_message)
        return jsonify({"error": error_message}), 500

# novels_db 示例 (Add the chapters field to the novel structure)
novels_db = {
    "novel_id_1": {
        "title": "我的第一部小说",
        "content": "...", # 可能废弃或作为默认/未分章节内容
        "characters": "...",
        "knowledge": "...",
        "style_prompt": "...",
        "character_tags": [],
        "glossary_tags": [],
        "style_tags": [],
        "chapters": []  # Add empty chapters list to store chapter data
    }
}

# Chapter Management API Endpoints
@app.route('/api/novels/<novel_id>/chapters', methods=['GET'])
def get_novel_chapters(novel_id):
    """Get all chapters for a specific novel"""
    if novel_id not in novels_db:
        return jsonify({"error": "Novel not found"}), 404
        
    # If no chapters field exists yet, initialize it
    if "chapters" not in novels_db[novel_id]:
        novels_db[novel_id]["chapters"] = []
        
    # Return chapters without content to reduce payload size
    chapters = []
    for chapter in novels_db[novel_id]["chapters"]:
        chapter_copy = {
            "id": chapter["id"],
            "title": chapter["title"],
            "order": chapter.get("order", 0),
            "created_at": chapter.get("created_at", "")
        }
        chapters.append(chapter_copy)
        
    return jsonify(chapters)

@app.route('/api/novels/<novel_id>/chapters', methods=['POST'])
def create_novel_chapter(novel_id):
    """Create a new chapter for a specific novel"""
    if novel_id not in novels_db:
        return jsonify({"error": "Novel not found"}), 404
        
    # Get request data
    data = request.json or {}
    title = data.get("title", "新章节")
    
    # Initialize chapters list if it doesn't exist
    if "chapters" not in novels_db[novel_id]:
        novels_db[novel_id]["chapters"] = []
        
    # Create a new chapter with UUID
    now = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    new_chapter = {
        "id": str(uuid.uuid4()),
        "title": title,
        "content": "",
        "order": len(novels_db[novel_id]["chapters"]),  # Add at the end
        "created_at": now
    }
    
    # Add the new chapter to the novel's chapters
    novels_db[novel_id]["chapters"].append(new_chapter)
    
    # Save to disk
    save_novels()
    
    # Return the new chapter (without content)
    chapter_copy = {key: val for key, val in new_chapter.items() if key != "content"}
    return jsonify(chapter_copy), 201

@app.route('/api/novels/<novel_id>/chapters/<chapter_id>', methods=['GET'])
def get_novel_chapter(novel_id, chapter_id):
    """Get a specific chapter with its content"""
    if novel_id not in novels_db:
        return jsonify({"error": "Novel not found"}), 404
        
    # Check if chapters field exists
    if "chapters" not in novels_db[novel_id]:
        return jsonify({"error": "Chapter not found"}), 404
        
    # Find the chapter by ID
    for chapter in novels_db[novel_id]["chapters"]:
        if chapter["id"] == chapter_id:
            return jsonify(chapter)
            
    return jsonify({"error": "Chapter not found"}), 404

@app.route('/api/novels/<novel_id>/chapters/<chapter_id>', methods=['PUT'])
def update_novel_chapter(novel_id, chapter_id):
    """Update a specific chapter"""
    if novel_id not in novels_db:
        return jsonify({"error": "Novel not found"}), 404
        
    # Check if chapters field exists
    if "chapters" not in novels_db[novel_id]:
        return jsonify({"error": "Chapter not found"}), 404
        
    # Get request data
    data = request.json or {}
    
    # Find and update the chapter
    for i, chapter in enumerate(novels_db[novel_id]["chapters"]):
        if chapter["id"] == chapter_id:
            # Update fields that were provided
            if "title" in data:
                novels_db[novel_id]["chapters"][i]["title"] = data["title"]
            if "content" in data:
                novels_db[novel_id]["chapters"][i]["content"] = data["content"]
            if "order" in data:
                novels_db[novel_id]["chapters"][i]["order"] = data["order"]
                
            # Save changes
            save_novels()
            
            # Return the updated chapter
            return jsonify(novels_db[novel_id]["chapters"][i])
            
    return jsonify({"error": "Chapter not found"}), 404

@app.route('/api/novels/<novel_id>/chapters/<chapter_id>', methods=['DELETE'])
def delete_novel_chapter(novel_id, chapter_id):
    """Delete a specific chapter"""
    if novel_id not in novels_db:
        return jsonify({"error": "Novel not found"}), 404
        
    # Check if chapters field exists
    if "chapters" not in novels_db[novel_id]:
        return jsonify({"error": "Chapter not found"}), 404
        
    # Find and remove the chapter
    for i, chapter in enumerate(novels_db[novel_id]["chapters"]):
        if chapter["id"] == chapter_id:
            # Remove the chapter
            removed_chapter = novels_db[novel_id]["chapters"].pop(i)
            
            # Reorder remaining chapters if needed
            for j, ch in enumerate(novels_db[novel_id]["chapters"]):
                ch["order"] = j
                
            # Save changes
            save_novels()
            
            return jsonify({"message": f"Chapter '{removed_chapter['title']}' deleted successfully"})
            
    return jsonify({"error": "Chapter not found"}), 404

@app.route('/api/novels/<novel_id>/chapters/reorder', methods=['PUT'])
def reorder_novel_chapters(novel_id):
    """Reorder chapters for a specific novel"""
    if novel_id not in novels_db:
        return jsonify({"error": "Novel not found"}), 404
        
    # Check if chapters field exists
    if "chapters" not in novels_db[novel_id]:
        novels_db[novel_id]["chapters"] = []
        return jsonify({"message": "No chapters to reorder"}), 200
        
    # Get request data - expecting an array of chapter IDs in the new order
    data = request.json or []
    
    if not isinstance(data, list):
        return jsonify({"error": "Invalid data format. Expected array of chapter IDs with order"}), 400
        
    # Create a map of chapter IDs to their current objects
    chapter_map = {chapter["id"]: chapter for chapter in novels_db[novel_id]["chapters"]}
    
    # Create a new ordered list of chapters
    new_chapters = []
    
    # Process each item in the request
    for i, item in enumerate(data):
        if isinstance(item, dict) and "id" in item:
            chapter_id = item["id"]
            if chapter_id in chapter_map:
                # Get the original chapter and update its order
                chapter = chapter_map[chapter_id].copy()
                chapter["order"] = i
                new_chapters.append(chapter)
                
    # If any chapters were missed in the request, append them at the end
    missed_chapters = [chapter for chapter_id, chapter in chapter_map.items() 
                      if chapter_id not in [item["id"] for item in data if isinstance(item, dict) and "id" in item]]
    
    for i, chapter in enumerate(missed_chapters):
        chapter["order"] = len(new_chapters) + i
        new_chapters.append(chapter)
        
    # Replace the chapters list
    novels_db[novel_id]["chapters"] = new_chapters
    
    # Save changes
    save_novels()
    
    return jsonify({"message": "Chapters reordered successfully"})

@app.route('/api/character-relationship-graph', methods=['POST'])
def generate_character_relationship_graph():
    """生成角色关系图谱，支持自定义API模型"""
    try:
        req = request.json or {}
        content = req.get('content', '')
        model = req.get('model', '')
        if not content:
            return jsonify({'error': '没有提供内容'}), 400
        if not model:
            return jsonify({'error': '未选择模型'}), 400

        # 优化后的prompt，更简洁且强制输出格式
        prompt = (
            "分析文本中的人物关系，输出JSON格式：\n"
            "{\"nodes\":[{\"id\":\"人物名\",\"name\":\"人物名\"}],\"links\":[{\"source\":\"人物1\",\"target\":\"人物2\",\"relation\":\"关系\"}]}\n"
            "要求：\n"
            "1. nodes包含所有人物，只需id和name字段\n"
            "2. links包含所有关系，必须有relation字段\n"
            "3. relation必须是简短中文词语\n"
            "4. 只输出JSON，不要其他内容\n"
            f"文本：{content}"
        )

        # 自定义API模型
        if model.startswith('custom-'):
            api_id = model.replace('custom-', '', 1)
            if api_id not in CUSTOM_API_CONFIGS:
                return jsonify({'error': '自定义API配置不存在'}), 404

            api_config = CUSTOM_API_CONFIGS[api_id]
            base_url = api_config['baseUrl']
            path = api_config.get('path', '')
            model_name = api_config['modelName']
            secret = api_config.get('secret', '')

            clean_base_url = base_url.rstrip('/')
            clean_path = path.lstrip('/')
            api_url = f"{clean_base_url}/{clean_path}" if clean_path else clean_base_url

            headers = {"Content-Type": "application/json"}
            if secret:
                headers["Authorization"] = f"Bearer {secret}"

            payload = {
                "model": model_name,
                "messages": [{"role": "user", "content": prompt}],
                "max_tokens": 2000,  # 增加token限制
                "temperature": 0.3    # 降低随机性
            }

            try:
                # 记录请求信息
                logging.info(f"发送请求到自定义API: {api_url}")
                logging.info(f"请求头: {headers}")
                logging.info(f"请求体: {json.dumps(payload, ensure_ascii=False)}")

                resp = requests.post(api_url, headers=headers, json=payload, timeout=60)
                
                # 记录响应信息
                logging.info(f"API响应状态码: {resp.status_code}")
                logging.info(f"API响应头: {dict(resp.headers)}")
                logging.info(f"API响应内容: {resp.text[:1000]}...")  # 只记录前1000个字符

                if resp.status_code != 200:
                    error_msg = f"自定义API分析失败，状态码: {resp.status_code}"
                    try:
                        error_detail = resp.json().get('error', {}).get('message', resp.text)
                        error_msg += f", 错误详情: {error_detail}"
                    except:
                        error_msg += f", 响应内容: {resp.text[:200]}..."
                    logging.error(error_msg)
                    return jsonify({'error': error_msg}), 500

                try:
                    result = resp.json()
                    logging.info(f"API响应JSON解析成功: {json.dumps(result, ensure_ascii=False)[:200]}...")

                    # 兼容OpenAI格式和常见格式
                    if 'choices' in result and result['choices']:
                        content = result['choices'][0].get('message',{}).get('content','')
                        if not content:
                            logging.error("API响应中没有content字段")
                            return jsonify({'error': 'API响应格式不正确：缺少content字段'}), 500
                            
                        content = strip_code_block(content)
                        logging.info(f"处理后的content: {content[:200]}...")
                        
                        # 检查JSON是否完整
                        if not content.strip().endswith('}'):
                            logging.error(f"API返回的JSON不完整: {content}")
                            return jsonify({'error': 'API返回的数据不完整，请尝试减少文本内容或使用其他模型'}), 500
                            
                        try:
                            relationship_data = json.loads(content)
                        except json.JSONDecodeError as e:
                            logging.error(f"JSON解析错误: {str(e)}, 原始内容: {content}")
                            return jsonify({'error': f'返回的数据格式不正确: {str(e)}'}), 500
                    else:
                        relationship_data = result

                    # 验证返回的数据结构
                    if not isinstance(relationship_data, dict):
                        raise ValueError("返回的数据不是字典格式")
                    if 'nodes' not in relationship_data:
                        raise ValueError("返回的数据缺少nodes字段")
                    if 'links' not in relationship_data:
                        raise ValueError("返回的数据缺少links字段")
                    if not isinstance(relationship_data['nodes'], list):
                        raise ValueError("nodes字段不是数组格式")
                    if not isinstance(relationship_data['links'], list):
                        raise ValueError("links字段不是数组格式")

                    return jsonify(relationship_data)

                except json.JSONDecodeError as e:
                    logging.error(f"JSON解析错误: {str(e)}, 原始内容: {resp.text[:1000]}...")
                    return jsonify({'error': f'返回的数据格式不正确: {str(e)}'}), 500
                except Exception as e:
                    logging.error(f"处理API响应时出错: {str(e)}, 原始内容: {resp.text[:1000]}...")
                    return jsonify({'error': f'处理API响应时出错: {str(e)}'}), 500

            except requests.exceptions.RequestException as e:
                logging.error(f"请求自定义API时出错: {str(e)}")
                return jsonify({'error': f'请求自定义API时出错: {str(e)}'}), 500
            except Exception as e:
                logging.error(f"自定义API调用出错: {str(e)}")
                return jsonify({'error': f'自定义API调用出错: {str(e)}'}), 500

        # OpenRouter模型
        else:
            if not OPENROUTER_API_KEY:
                return jsonify({'error': '请先设置OpenRouter API密钥'}), 401

            try:
                response = requests.post(
                    f"{OPENROUTER_API_BASE}/chat/completions",
                    headers={
                        "Authorization": f"Bearer {OPENROUTER_API_KEY}",
                        "Content-Type": "application/json"
                    },
                    json={
                        "model": model,
                        "messages": [{"role": "user", "content": prompt}],
                        "max_tokens": 2000,  # 增加token限制
                        "temperature": 0.3    # 降低随机性
                    }
                )

                # 记录响应信息
                logging.info(f"OpenRouter响应状态码: {response.status_code}")
                logging.info(f"OpenRouter响应内容: {response.text[:1000]}...")

                if response.status_code != 200:
                    error_msg = f"AI分析失败，状态码: {response.status_code}"
                    try:
                        error_detail = response.json().get('error', {}).get('message', response.text)
                        error_msg += f", 错误详情: {error_detail}"
                    except:
                        error_msg += f", 响应内容: {response.text[:200]}..."
                    logging.error(error_msg)
                    return jsonify({'error': error_msg}), 500

                try:
                    result = response.json()
                    content = result['choices'][0]['message']['content']
                    if not content:
                        logging.error("OpenRouter响应中没有content字段")
                        return jsonify({'error': 'OpenRouter响应格式不正确：缺少content字段'}), 500
                        
                    content = strip_code_block(content)
                    logging.info(f"处理后的content: {content[:200]}...")
                    
                    # 检查JSON是否完整
                    if not content.strip().endswith('}'):
                        logging.error(f"API返回的JSON不完整: {content}")
                        return jsonify({'error': 'API返回的数据不完整，请尝试减少文本内容或使用其他模型'}), 500
                        
                    relationship_data = json.loads(content)
                    
                    # 验证返回的数据结构
                    if not isinstance(relationship_data, dict):
                        raise ValueError("返回的数据不是字典格式")
                    if 'nodes' not in relationship_data:
                        raise ValueError("返回的数据缺少nodes字段")
                    if 'links' not in relationship_data:
                        raise ValueError("返回的数据缺少links字段")
                    if not isinstance(relationship_data['nodes'], list):
                        raise ValueError("nodes字段不是数组格式")
                    if not isinstance(relationship_data['links'], list):
                        raise ValueError("links字段不是数组格式")
                        
                    return jsonify(relationship_data)

                except json.JSONDecodeError as e:
                    logging.error(f"JSON解析错误: {str(e)}, 原始内容: {content}")
                    return jsonify({'error': f'返回的数据格式不正确: {str(e)}'}), 500
                except Exception as e:
                    logging.error(f"处理OpenRouter响应时出错: {str(e)}, 原始内容: {response.text[:1000]}...")
                    return jsonify({'error': f'处理OpenRouter响应时出错: {str(e)}'}), 500

            except requests.exceptions.RequestException as e:
                logging.error(f"请求OpenRouter时出错: {str(e)}")
                return jsonify({'error': f'请求OpenRouter时出错: {str(e)}'}), 500
            except Exception as e:
                logging.error(f"OpenRouter调用出错: {str(e)}")
                return jsonify({'error': f'OpenRouter调用出错: {str(e)}'}), 500

    except Exception as e:
        logging.error(f"生成角色关系图谱时出错: {str(e)}")
        return jsonify({'error': str(e)}), 500

def strip_code_block(text):
    """
    去除AI返回内容中的```json、```text等代码块包裹。
    """
    import re
    # 匹配```json ... ```或```text ... ```或``` ... ```
    pattern = r"^```(?:json|text)?\s*([\s\S]*?)\s*```$"
    match = re.match(pattern, text.strip(), re.IGNORECASE)
    if match:
        return match.group(1).strip()
    return text.strip()

@app.route('/api/knowledge-graph', methods=['POST'])
def generate_knowledge_graph():
    """生成知识图谱，支持自定义API模型"""
    try:
        req = request.json or {}
        content = req.get('content', '')
        model = req.get('model', '')
        if not content:
            return jsonify({'error': '没有提供内容'}), 400
        if not model:
            return jsonify({'error': '未选择模型'}), 400

        # 优化后的prompt，更简洁且强制输出格式
        prompt = (
            "分析文本中的知识点和关系，输出JSON格式：\n"
            "{\"nodes\":[{\"id\":\"知识点ID\",\"name\":\"知识点名称\",\"category\":\"分类\",\"tags\":[\"标签1\",\"标签2\"]}],"
            "\"links\":[{\"source\":\"知识点1\",\"target\":\"知识点2\",\"relation\":\"关系\",\"type\":\"关系类型\"}]}\n"
            "要求：\n"
            "1. nodes包含所有知识点，必须有id、name、category和tags字段\n"
            "2. category必须是以下之一：人物、地点、物品、概念、事件、规则\n"
            "3. tags是相关标签数组，至少包含一个标签\n"
            "4. links包含所有关系，必须有source、target、relation和type字段\n"
            "5. type必须是以下之一：包含、属于、导致、影响、引用\n"
            "6. relation必须是简短中文词语\n"
            "7. 只输出JSON，不要其他内容\n"
            f"文本：{content}"
        )

        # 自定义API模型
        if model.startswith('custom-'):
            api_id = model.replace('custom-', '', 1)
            if api_id not in CUSTOM_API_CONFIGS:
                return jsonify({'error': '自定义API配置不存在'}), 404

            api_config = CUSTOM_API_CONFIGS[api_id]
            base_url = api_config['baseUrl']
            path = api_config.get('path', '')
            model_name = api_config['modelName']
            secret = api_config.get('secret', '')

            clean_base_url = base_url.rstrip('/')
            clean_path = path.lstrip('/')
            api_url = f"{clean_base_url}/{clean_path}" if clean_path else clean_base_url

            headers = {"Content-Type": "application/json"}
            if secret:
                headers["Authorization"] = f"Bearer {secret}"

            payload = {
                "model": model_name,
                "messages": [{"role": "user", "content": prompt}],
                "max_tokens": 2000,
                "temperature": 0.3
            }

            try:
                resp = requests.post(api_url, headers=headers, json=payload, timeout=60)
                
                if resp.status_code != 200:
                    error_msg = f"自定义API分析失败，状态码: {resp.status_code}"
                    try:
                        error_detail = resp.json().get('error', {}).get('message', resp.text)
                        error_msg += f", 错误详情: {error_detail}"
                    except:
                        error_msg += f", 响应内容: {resp.text[:200]}..."
                    logging.error(error_msg)
                    return jsonify({'error': error_msg}), 500

                try:
                    result = resp.json()
                    content = result['choices'][0]['message']['content']
                    if not content:
                        return jsonify({'error': 'API响应格式不正确：缺少content字段'}), 500
                        
                    content = strip_code_block(content)
                    
                    if not content.strip().endswith('}'):
                        return jsonify({'error': 'API返回的数据不完整，请尝试减少文本内容或使用其他模型'}), 500
                        
                    knowledge_data = json.loads(content)
                    
                    # 验证返回的数据结构
                    if not isinstance(knowledge_data, dict):
                        raise ValueError("返回的数据不是字典格式")
                    if 'nodes' not in knowledge_data:
                        raise ValueError("返回的数据缺少nodes字段")
                    if 'links' not in knowledge_data:
                        raise ValueError("返回的数据缺少links字段")
                    if not isinstance(knowledge_data['nodes'], list):
                        raise ValueError("nodes字段不是数组格式")
                    if not isinstance(knowledge_data['links'], list):
                        raise ValueError("links字段不是数组格式")
                        
                    return jsonify(knowledge_data)

                except json.JSONDecodeError as e:
                    return jsonify({'error': f'返回的数据格式不正确: {str(e)}'}), 500
                except Exception as e:
                    return jsonify({'error': f'处理API响应时出错: {str(e)}'}), 500

            except requests.exceptions.RequestException as e:
                return jsonify({'error': f'请求自定义API时出错: {str(e)}'}), 500
            except Exception as e:
                return jsonify({'error': f'自定义API调用出错: {str(e)}'}), 500

        # OpenRouter模型
        else:
            if not OPENROUTER_API_KEY:
                return jsonify({'error': '请先设置OpenRouter API密钥'}), 401

            try:
                response = requests.post(
                    f"{OPENROUTER_API_BASE}/chat/completions",
                    headers={
                        "Authorization": f"Bearer {OPENROUTER_API_KEY}",
                        "Content-Type": "application/json"
                    },
                    json={
                        "model": model,
                        "messages": [{"role": "user", "content": prompt}],
                        "max_tokens": 2000,
                        "temperature": 0.3
                    }
                )

                if response.status_code != 200:
                    error_msg = f"AI分析失败，状态码: {response.status_code}"
                    try:
                        error_detail = response.json().get('error', {}).get('message', response.text)
                        error_msg += f", 错误详情: {error_detail}"
                    except:
                        error_msg += f", 响应内容: {response.text[:200]}..."
                    return jsonify({'error': error_msg}), 500

                try:
                    result = response.json()
                    content = result['choices'][0]['message']['content']
                    if not content:
                        return jsonify({'error': 'OpenRouter响应格式不正确：缺少content字段'}), 500
                        
                    content = strip_code_block(content)
                    
                    if not content.strip().endswith('}'):
                        return jsonify({'error': 'API返回的数据不完整，请尝试减少文本内容或使用其他模型'}), 500
                        
                    knowledge_data = json.loads(content)
                    
                    # 验证返回的数据结构
                    if not isinstance(knowledge_data, dict):
                        raise ValueError("返回的数据不是字典格式")
                    if 'nodes' not in knowledge_data:
                        raise ValueError("返回的数据缺少nodes字段")
                    if 'links' not in knowledge_data:
                        raise ValueError("返回的数据缺少links字段")
                    if not isinstance(knowledge_data['nodes'], list):
                        raise ValueError("nodes字段不是数组格式")
                    if not isinstance(knowledge_data['links'], list):
                        raise ValueError("links字段不是数组格式")
                        
                    return jsonify(knowledge_data)

                except json.JSONDecodeError as e:
                    return jsonify({'error': f'返回的数据格式不正确: {str(e)}'}), 500
                except Exception as e:
                    return jsonify({'error': f'处理OpenRouter响应时出错: {str(e)}'}), 500

            except requests.exceptions.RequestException as e:
                return jsonify({'error': f'请求OpenRouter时出错: {str(e)}'}), 500
            except Exception as e:
                return jsonify({'error': f'OpenRouter调用出错: {str(e)}'}), 500

    except Exception as e:
        return jsonify({'error': str(e)}), 500

if __name__ == '__main__':
    # 确保静态文件夹配置正确
    import webbrowser
    from threading import Timer
    
    # 设置端口
    port = 5000
    
    # 打包后的环境自动启用生产模式
    debug_mode = not getattr(sys, 'frozen', False)
    
    # 添加标志变量防止多次打开浏览器
    browser_opened = False
    
    # 定义一个函数，在应用启动后打开浏览器
    def open_browser():
        global browser_opened
        # 仅在browser_opened为False时打开浏览器
        if not browser_opened:
            webbrowser.open(f'http://127.0.0.1:{port}/')
            browser_opened = True
            print("笔墨星河已成功启动！浏览器窗口即将打开。")
            print(f"如果没有自动打开，请手动访问: http://127.0.0.1:{port}/")
            print("笔墨星河，本软件免费，请合法合规使用本软件，请勿用作非法用途。软件作者：L.H")    
    # 在新线程中启动浏览器（延迟1.5秒，确保Flask服务器已启动）
    # 增加延迟时间，确保服务器完全启动
    Timer(1.5, open_browser).start()
    
    print("正在启动笔墨星河应用...")
    # 在启动应用前加载数据
    load_data()
    # 启动Flask应用，禁用Flask的自动重载器，这可能导致浏览器被打开两次
    app.run(debug=debug_mode, port=port, use_reloader=False)